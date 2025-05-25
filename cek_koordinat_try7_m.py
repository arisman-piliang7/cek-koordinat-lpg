import streamlit as st
import pandas as pd
import math
import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from zipfile import ZipFile
import re

def haversine(lat1, lon1, lat2, lon2):
    R = 6371.0  # Radius bumi dalam km
    phi1, phi2 = math.radians(lat1), math.radians(lat2)
    dphi = math.radians(lat2 - lat1)
    dlambda = math.radians(lon2 - lon1)
    a = math.sin(dphi / 2)**2 + math.cos(phi1) * math.cos(phi2) * math.sin(dlambda / 2)**2
    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))
    return R * c  # hasil km

def is_valid_coordinate(coord):
    if pd.isna(coord):
        return False, "Nilai kosong"
    if isinstance(coord, str):
        coord = coord.strip().replace(",", ".").replace('"', "").replace("'", "").upper()
        if coord in ("", "NULL", "NA", "N/A", "NONE", "-"):
            return False, "Nilai kosong atau tidak valid"
        if re.search(r"[^0-9.\-]", coord):
            return False, "Mengandung karakter tidak valid (spasi/tanda baca)"
        try:
            float(coord)
            return True, ""
        except:
            return False, "Format angka tidak valid"
    try:
        float(coord)
        return True, ""
    except:
        return False, "Format angka tidak valid"

def clean_coordinate(coord):
    try:
        if pd.isna(coord):
            return None
        if isinstance(coord, str):
            coord = coord.strip().replace(",", ".").replace('"', "").replace("'", "").upper()
            if coord in ("", "NULL", "NA", "N/A", "NONE", "-"):
                return None
            coord = re.sub(r"[^0-9.\-]", "", coord)
        return float(coord)
    except:
        return None

st.title("Evaluasi Jarak Koordinat Pangkalan LPG 3 Kg")

uploaded_file = st.file_uploader("Unggah file CSV format", type=["csv"])

# Inisialisasi session_state
if "last_uploaded_filename" not in st.session_state:
    st.session_state["last_uploaded_filename"] = None
if "koordinat_bersih" not in st.session_state:
    st.session_state["koordinat_bersih"] = False
if "hasil_df" not in st.session_state:
    st.session_state["hasil_df"] = None
if "word_files" not in st.session_state:
    st.session_state["word_files"] = []

if uploaded_file is not None:
    # Reset session state jika file berbeda
    if uploaded_file.name != st.session_state["last_uploaded_filename"]:
        for key in list(st.session_state.keys()):
            if key not in ("last_uploaded_filename", "koordinat_bersih", "hasil_df", "word_files"):
                del st.session_state[key]
        st.session_state["koordinat_bersih"] = False
        st.session_state["hasil_df"] = None
        st.session_state["word_files"] = []
        st.session_state["last_uploaded_filename"] = uploaded_file.name

    df = pd.read_csv(uploaded_file)
    st.write("Data Awal:")
    st.dataframe(df)

    # Indeks kolom sesuai file CSV
    lat_index = 8
    lon_index = 9
    soldtoparty_index = 0
    nama_agen_index = 1
    nama_pangkalan_index = 2

    # Validasi koordinat awal
    invalid_rows = []
    for idx, row in df.iterrows():
        lat = row[lat_index]
        lon = row[lon_index]
        valid_lat, reason_lat = is_valid_coordinate(lat)
        valid_lon, reason_lon = is_valid_coordinate(lon)
        if not valid_lat or not valid_lon:
            reason = reason_lat if not valid_lat else reason_lon
            invalid_rows.append((idx + 2, row[nama_pangkalan_index], row[nama_agen_index], reason))

    # Jika koordinat belum dibersihkan
    if not st.session_state["koordinat_bersih"]:
        if invalid_rows:
            jumlah_invalid = len(invalid_rows)
            st.warning(f"Terdapat koordinat yang tidak valid sejumlah {jumlah_invalid} baris:")

            # Tampilkan daftar error rapi
            for baris, pangkalan, agen, reason in invalid_rows:
                st.write(f"- Baris ke-{baris}, Pangkalan: {pangkalan}, Agen: {agen} — Alasan: {reason}")

            if st.button("PERBAIKI OTOMATIS"):
                gagal_diperbaiki = []
                for i, row in df.iterrows():
                    lat = clean_coordinate(row[lat_index])
                    lon = clean_coordinate(row[lon_index])
                    if lat is not None and lon is not None:
                        df.at[i, df.columns[lat_index]] = lat
                        df.at[i, df.columns[lon_index]] = lon
                    else:
                        gagal_diperbaiki.append((i + 2, row[nama_pangkalan_index], row[nama_agen_index]))

                if gagal_diperbaiki:
                    st.error("Beberapa data tidak dapat diperbaiki secara otomatis:")
                    for baris, pangkalan, agen in gagal_diperbaiki:
                        st.write(f"- Baris ke-{baris}, Pangkalan: {pangkalan}, Agen: {agen}")
                    st.warning("Silakan perbaiki koordinat secara manual dan unggah ulang file CSV-nya.")
                    st.stop()
                else:
                    st.success("Semua koordinat berhasil diperbaiki secara otomatis.")
                    st.session_state["koordinat_bersih"] = True
                    st.success("Silakan tentukan jarak minimal pangkalan dan jumlah jarak kemudian tekan tombol 'PROSES VALIDASI' untuk melanjutkan.")
            else:
                st.stop()
        else:
            st.success("Semua koordinat sudah valid.")
            st.session_state["koordinat_bersih"] = True

    # Jika koordinat sudah bersih, tampilkan form validasi jarak
    if st.session_state["koordinat_bersih"]:
        with st.form("validasi_form"):
            batas_meter = st.slider("Pilih batas jarak antar Pangkalan (meter):", 10, 1000, 100)
            batas_km = batas_meter / 1000

            grouped = df.groupby(df.columns[soldtoparty_index])
            max_length = max(len(group) for _, group in grouped)
            max_slider = max_length - 1 if max_length > 1 else 1

            slider_max = st.slider(
                "Jumlah kolom Jarak yang ingin ditampilkan:",
                min_value=1,
                max_value=max_slider,
                value=min(10, max_slider)
            )

            submit = st.form_submit_button("PROSES VALIDASI")

        if submit:
            word_files = []
            all_group_dfs = []

            for soldtoparty, group in grouped:
                group = group.reset_index(drop=True)
                nama_agen = group.iloc[0, nama_agen_index]
                koordinat = [
                    (
                        clean_coordinate(row[lat_index]) or 0.0,
                        clean_coordinate(row[lon_index]) or 0.0
                    )
                    for _, row in group.iterrows()
                ]
                n = len(koordinat)

                # Hitung jarak antar koordinat untuk kolom Jarak 1..slider_max (dalam meter)
                for d in range(1, slider_max + 1):
                    jarak_list = []
                    for i in range(n):
                        if i >= d:
                            lat1, lon1 = koordinat[i - d]
                            lat2, lon2 = koordinat[i]
                            jarak_km = haversine(lat1, lon1, lat2, lon2)
                            jarak_meter = round(jarak_km * 1000, 2)  # Konversi ke meter
                            jarak_list.append(jarak_meter)
                        else:
                            jarak_list.append("")
                    group[f'Jarak {d} (m)'] = jarak_list

                all_group_dfs.append(group)

                # Buat rekap jarak dibawah batas untuk surat dokumen (meter)
                rekap_bawah = []
                for d in range(1, slider_max + 1):
                    for i in range(n):
                        if i >= d:
                            jarak = group.loc[i, f'Jarak {d} (m)']
                            if jarak != "" and jarak < batas_meter:
                                pangkalan_1 = group.loc[i - d, group.columns[nama_pangkalan_index]]
                                pangkalan_2 = group.loc[i, group.columns[nama_pangkalan_index]]
                                rekap_bawah.append({"pangkalan_1": pangkalan_1, "pangkalan_2": pangkalan_2, "jarak": jarak})

                # Jika ada data jarak dibawah batas buat dokumen word
                if rekap_bawah:
                    doc = Document()
                    style = doc.styles['Normal']
                    font = style.font
                    font.name = 'Arial'
                    font.size = Pt(12)

                    def add_paragraph_justify(text):
                        p = doc.add_paragraph(text)
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p.paragraph_format.space_after = Pt(0)

                    doc.add_paragraph("Medan, Januari 2025").alignment = WD_ALIGN_PARAGRAPH.LEFT
                    doc.add_paragraph("No. /XXXXXXXXX/2025-XX").alignment = WD_ALIGN_PARAGRAPH.LEFT
                    doc.add_paragraph("Lampiran:")

                    perihal_paragraph = doc.add_paragraph()
                    run_perihal = perihal_paragraph.add_run(f"Perihal: Evaluasi Data Pangkalan {nama_agen} pada SIMELON")
                    run_perihal.bold = True
                    perihal_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    perihal_paragraph.paragraph_format.space_after = Pt(0)

                    doc.add_paragraph("Yang terhormat\nPimpinan di Tempat")

                    add_paragraph_justify("\nDengan hormat,")
                    add_paragraph_justify("\nDalam rangka menjamin kemudahan akses masyarakat untuk mendapatkan LPG 3 Kg, maka kami telah melakukan evaluasi data lokasi pangkalan dari data SIMELON Agen LPG 3 Kg Saudara.")
                    add_paragraph_justify(f"\nHasil evaluasi tersebut ditemukan bahwa terdapat pangkalan dengan titik lokasi (latitude, longitude) dibawah {batas_meter} meter yaitu:")
                    for idx, row in enumerate(rekap_bawah):
                        add_paragraph_justify(f"{idx+1}. Pangkalan {row['pangkalan_1']} dengan Pangkalan {row['pangkalan_2']}, jarak {row['jarak']} meter")
                    add_paragraph_justify("\nSehubungan dengan hal tersebut, maka kami minta Saudara melakukan evaluasi berupa:")
                    add_paragraph_justify("1. Memastikan kembali titik lokasi pangkalan sesuai dengan kondisi riil lapangan dan mengupdate pada Web SIMELON.")
                    add_paragraph_justify("2. Apabila pangkalan benar pada titik lokasi yang sama, maka segera lakukan pemindahan lokasi salah satu pangkalan.")
                    add_paragraph_justify("\nSelanjutnya agar Saudara segera menindaklanjuti temuan tersebut dan melaporkan kembali kepada kami dalam waktu 1 bulan kedepan.")
                    add_paragraph_justify("\nDemikian disampaikan, atas perhatian dan kerjasamanya kami ucapkan terima kasih.")
                    doc.add_paragraph("\nJabatan Manager")
                    doc.add_paragraph("Nama Manager")
                    doc.add_paragraph("Tembusan:")
                    doc.add_paragraph("1. Executive GM Regional Sumbagut")
                    doc.add_paragraph("2. SAM Retail Terkait")

                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    filename = f"Evaluasi Data Pangkalan {nama_agen}.docx"
                    word_files.append((filename, buffer.read()))

            # Gabungkan semua group dataframe jadi satu
            st.session_state["hasil_df"] = pd.concat(all_group_dfs, ignore_index=True)
            st.session_state["word_files"] = word_files

            # Simpan hasil ke Excel dengan highlight jarak < batas (meter)
            excel_buffer = io.BytesIO()
            with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
                df_final = st.session_state["hasil_df"]
                df_final.to_excel(writer, index=False, sheet_name='Hasil Validasi')
                workbook = writer.book
                worksheet = writer.sheets['Hasil Validasi']

                # Format highlight: stabilo kuning + font merah
                format_highlight = workbook.add_format({'font_color': 'red', 'bg_color': '#FFFF00'})

                jarak_cols = [col for col in df_final.columns if col.startswith("Jarak ")]
                for col_index, col_name in enumerate(df_final.columns):
                    if col_name in jarak_cols:
                        for row_idx, value in enumerate(df_final[col_name]):
                            if isinstance(value, (int, float)) and value < batas_meter:
                                worksheet.write(row_idx + 1, col_index, value, format_highlight)

            excel_buffer.seek(0)
            st.download_button(
                "Unduh hasil_jarak_format.xlsx",
                data=excel_buffer,
                file_name="hasil_jarak_format.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

if st.session_state.get("word_files"):
    zip_buffer = io.BytesIO()
    with ZipFile(zip_buffer, "w") as zip_file:
        for filename, data in st.session_state["word_files"]:
            zip_file.writestr(filename, data)
    zip_buffer.seek(0)
    st.download_button("Unduh Semua Rekap Agen (ZIP)", data=zip_buffer, file_name="rekap_agen.zip", mime="application/zip")
